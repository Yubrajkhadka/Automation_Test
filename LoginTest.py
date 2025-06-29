import pytest
import os
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from docx import Document
import time

@pytest.fixture
def setup():
    options = webdriver.EdgeOptions()
    options.add_argument("--start-maximized")  # Optional for better visibility
    driver = webdriver.Edge(service=Service("msedgedriver.exe"), options=options)
    yield driver
    driver.quit()

@pytest.fixture(scope="session", autouse=True)
def clean_log_file():
    log_dir = "logs"
    doc_path = os.path.join(log_dir, "LoginTest_SuccessLog.docx")
    if os.path.exists(doc_path):
        try:
            os.remove(doc_path)
            print("Old log file deleted before tests")
        except Exception as e:
            print(f"Warning: Could not delete old log file: {e}")
    yield

def log_success_to_word(status, username, password, client_code):
    from docx import Document
    from datetime import datetime
    import os

    log_dir = "logs"
    os.makedirs(log_dir, exist_ok=True)
    doc_path = os.path.join(log_dir, "LoginTest_Case.docx")

    try:
        if os.path.exists(doc_path):
            doc = Document(doc_path)
            if doc.tables:
                table = doc.tables[0]
            else:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Status'
                hdr_cells[1].text = 'Username'
                hdr_cells[2].text = 'Password'
                hdr_cells[3].text = 'Client Code'
                hdr_cells[4].text = 'Timestamp'
        else:
            doc = Document()
            doc.add_heading("Login Test Cases", 0)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Status'
            hdr_cells[1].text = 'Username'
            hdr_cells[2].text = 'Password'
            hdr_cells[3].text = 'Client Code'
            hdr_cells[4].text = 'Timestamp'

        row_cells = table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = username
        row_cells[2].text = password
        row_cells[3].text = client_code
        row_cells[4].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        doc.save(doc_path)

    except PermissionError:
        print(f"❌ Could not save to {doc_path}. Please close the file if it's open.")

@pytest.mark.parametrize("client_code, username, password", [
    ("018", "asmita.kafle", "Admin@123")

])
def test_login_success(setup,client_code,username,password):
    driver = setup
    wait = WebDriverWait(driver, 15)
    driver.get("https://vaptcrs.infodev.com.np/login-client")
    # Fill form
    wait.until(EC.element_to_be_clickable((By.NAME, "clientCode"))).send_keys(client_code + Keys.TAB)
    wait.until(EC.element_to_be_clickable((By.NAME, "username"))).send_keys(username)
    wait.until(EC.element_to_be_clickable((By.NAME, "password"))).send_keys(password)
    wait.until(EC.element_to_be_clickable((By.XPATH, "//button[@type='submit']"))).click()

    # Wait for redirect or any post-login element (adjust selector as needed)
    try:
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "a[data-tooltip-id='Dashboard']")))
        log_success_to_word("Pass",client_code,username,password)
        create_new_user(driver,wait)
    except:
        log_success_to_word("Fail", client_code,username,password)
        pytest.fail("Login failed or dashboard not loaded.")

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pytest
import time

def create_new_user(driver, wait: WebDriverWait):
    try:
        # Navigate to User Management
        wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'a[data-tooltip-id="User Management"]'))).click()
        
        # Click Add button
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Add']"))).click()
        
        # Fill in user details
        wait.until(EC.element_to_be_clickable((By.NAME, "username"))).send_keys("Yubraj11")
        wait.until(EC.element_to_be_clickable((By.NAME, "firstName"))).send_keys("Yubraj")
        wait.until(EC.element_to_be_clickable((By.NAME, "middleName"))).send_keys("")
        wait.until(EC.element_to_be_clickable((By.NAME, "lastName"))).send_keys("Khadka")
        wait.until(EC.element_to_be_clickable((By.NAME, "email"))).send_keys("yubraj.khadka@infodev.com.np")
        wait.until(EC.element_to_be_clickable((By.NAME, "phoneNumber"))).send_keys("9861562381")
        wait.until(EC.element_to_be_clickable((By.NAME, "section"))).send_keys("IT")
        
        # Select Branch
        branch_dropdown = wait.until(EC.element_to_be_clickable((
            By.XPATH, "//label[contains(text(), 'Branch')]/following::div[contains(@class, 'css-9jkw19')]"
        )))
        branch_dropdown.click()
        branch_input = wait.until(EC.presence_of_element_located((
            By.CSS_SELECTOR, "input[id^='react-select'][type='text']"
        )))
        branch_input.send_keys("Head Office")
        time.sleep(1)  # small pause for dropdown to populate
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Head Office']"))).click()
        time.sleep(1)
        
        # Select Department
        department_dropdown = wait.until(EC.element_to_be_clickable((
            By.XPATH, "//label[contains(text(), 'Department')]/following::div[contains(@class, 'css-9jkw19')]"
        )))
        department_dropdown.click()
        department_input = department_dropdown.find_element(By.XPATH, ".//input")
        department_input.send_keys("Head Department")
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Head Department']"))).click()
        time.sleep(1)
        
        # Select User Type
        user_type_dropdown = wait.until(EC.element_to_be_clickable((
            By.XPATH, "//label[contains(text(), 'User Type')]/following::div[contains(@class, 'css-9jkw19')]"
        )))
        user_type_dropdown.click()
        user_type_input = user_type_dropdown.find_element(By.XPATH, ".//input")
        user_type_input.send_keys("Branch Admin")
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Branch Admin']"))).click()
        time.sleep(1)
        
        # Select Role (fixed typo and XPath)
        role_dropdown = wait.until(EC.element_to_be_clickable((
            By.XPATH, "//label[contains(text(),'Roles')]/following::div[contains(@class,'css-5fiuwh')]"
        )))
        role_dropdown.click()
        role_input = role_dropdown.find_element(By.XPATH, ".//input")
        role_input.send_keys("Nepal Admin")
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Nepal Admin']"))).click()
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//div[@class='w-full' and text()='Save']"))).click()
        time.sleep(4)


    except Exception as e:
        print(f"❌ Failed in user creation flow: {e}")
        pytest.fail("Could not complete new user creation")
